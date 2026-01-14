import streamlit as st
import pandas as pd
import io
import time
import zipfile
import re

# ==============================================================================
# 依赖库检查
# ==============================================================================
try:
    import docx
    from docx.shared import Pt, Cm, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from openpyxl.styles import Border, Side, Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

# ==============================================================================
# Zone 0: 全局配置 & 样式注入 (核心修复)
# ==============================================================================
st.set_page_config(page_title="淡藤财务报表 Pro", page_icon="😈", layout="wide", initial_sidebar_state="expanded")

if "prank" in st.query_params:
    st.session_state.page = 'mapping'
    st.session_state.prank_solved = True
    st.query_params.clear()

def inject_css():
    st.markdown("""
    <style>
        :root { --bg-color: #0d1117; --card-bg: #161b22; --text: #c9d1d9; --border-color: #30363d; }
        .stApp { background-color: var(--bg-color); color: var(--text); }
        button p { white-space: nowrap !important; }

        .file-name { font-weight: 600; font-size: 14px; color: #e6edf3; display: block; line-height: 1.2; }
        .file-stats { font-size: 12px; color: #8b949e; display: block; margin-top: 2px; }
        .file-icon { font-size: 24px; display: flex; align-items: center; justify-content: center; height: 100%; }
        
        div[data-testid="column"] { overflow: visible !important; }
        div[data-testid="column"] button[kind="secondary"] {
            border: 1px solid rgba(255,255,255,0.1) !important;
            background-color: rgba(255,255,255,0.05) !important;
            color: #c9d1d9 !important;
            padding: 0px !important; margin: 0px !important;
            height: 42px !important; width: 100% !important; min-width: 40px !important;
            display: flex !important; align-items: center !important; justify-content: center !important;
            float: none !important; border-radius: 6px !important; transition: all 0.2s;
        }
        div[data-testid="column"] button[kind="secondary"]:hover {
            color: #ff7b72 !important; border-color: #ff7b72 !important;
            background-color: rgba(255, 123, 114, 0.1) !important;
        }
        
        /* ==========================================================================
           🚀 核心布局修复：侧边栏底部固定方案
           ========================================================================== */
        
        /* 1. 强制侧边栏滚动区域撑满屏幕高度，并开启 Flex 布局 */
        [data-testid="stSidebarUserContent"] > div:first-child {
            height: 99vh;  /* 稍微留一点余量防止双滚动条 */
            display: flex;
            flex-direction: column;
        }

        /* 2. 针对 Tertiary 按钮 (猫咪图标) 的特殊样式 */
        button[kind="tertiary"] {
            border: none !important; 
            background: transparent !important; 
            box-shadow: none !important;
            padding: 0 !important;
            width: 100% !important; 
            margin: 0 !important;
        }

        /* 3. 强制让按钮内部的 Emoji 文本居中 */
        button[kind="tertiary"] > div[data-testid="stMarkdownContainer"] {
            width: 100%;
            display: flex;
            justify-content: center !important; /* 水平居中 */
            align-items: center !important;
        }
        
        button[kind="tertiary"] > div[data-testid="stMarkdownContainer"] > p {
            font-size: 3rem !important; /* 放大图标 */
            margin: 0 !important;
            padding: 0 !important;
            text-align: center !important;
        }

        button[kind="tertiary"]:hover { 
            transform: scale(1.25) rotate(5deg); 
            background: transparent !important;
            color: #58a6ff !important;
            transition: transform 0.2s cubic-bezier(0.175, 0.885, 0.32, 1.275);
        }
        
        /* ========================================================================== */

        .sidebar-section { margin-bottom: 20px; }
        .sidebar-label { font-size: 0.85rem; color: #8b949e; margin-bottom: 5px; font-weight: 600; }
        .nav-header { font-size: 1.2rem; font-weight: bold; display:flex; align-items:center; height: 100%; }
        .info-bar { background-color: rgba(56, 139, 253, 0.1); border-left: 4px solid #58a6ff; color: #c9d1d9; padding: 8px 15px; margin-bottom: 20px; font-size: 0.9rem; border-radius: 4px; }
        .error-box { border: 1px solid #ff7b72; background-color: rgba(255, 123, 114, 0.1); padding: 15px; border-radius: 6px; margin-bottom: 15px; }
        .balance-box-ok { border: 1px solid #238636; background-color: rgba(35, 134, 54, 0.1); padding: 10px; border-radius: 6px; margin-bottom: 15px; color: #3fb950; }
        .balance-box-err { border: 1px solid #da3633; background-color: rgba(218, 54, 51, 0.1); padding: 10px; border-radius: 6px; margin-bottom: 15px; color: #f85149; font-weight: bold;}

        /* --- 强制让 Selectbox 内容居中 --- */
        div[data-baseweb="select"] > div {
            justify-content: center !important;
            text-align: center !important;
        }
        div[data-baseweb="select"] span {
            width: 100%;
            text-align: center;
        }
    </style>
    """, unsafe_allow_html=True)

# ==============================================================================
# Zone A: 纯逻辑层
# ==============================================================================
class TemplateManager:
    DEFAULT_NAME = "默认模板"

    @staticmethod
    def init_defaults():
        if 'templates' not in st.session_state:
            st.session_state.templates = {
                TemplateManager.DEFAULT_NAME: DataEngine.get_default_config()
            }
        if 'active_template_name' not in st.session_state:
            st.session_state.active_template_name = TemplateManager.DEFAULT_NAME
        if 'editing_template_name' not in st.session_state:
            st.session_state.editing_template_name = TemplateManager.DEFAULT_NAME
        # Ensure params exist
        if 'params' not in st.session_state:
             st.session_state.params = {
                'price': 1600, 
                'hours_limit': 100, 
                'sub_tag': '差旅补助', 
                'period': '2025Q1'
             }

    @staticmethod
    def get_template(name):
        return st.session_state.templates.get(name, DataEngine.get_default_config())

    @staticmethod
    def save_template(name, df_config):
        st.session_state.templates[name] = df_config

    @staticmethod
    def delete_template(name):
        if name in st.session_state.templates and name != TemplateManager.DEFAULT_NAME:
            del st.session_state.templates[name]
            if st.session_state.active_template_name == name:
                st.session_state.active_template_name = TemplateManager.DEFAULT_NAME
            if st.session_state.editing_template_name == name:
                st.session_state.editing_template_name = TemplateManager.DEFAULT_NAME
            return True
        return False

    @staticmethod
    def get_all_names():
        return list(st.session_state.templates.keys())

class DataEngine:
    @staticmethod
    def get_quarter_str(text):
        if not text: return "2025Qx"
        s = str(text)
        year_match = re.search(r'(\d{4})', s)
        year = year_match.group(1) if year_match else "2025"
        quarter = "Q1" 
        if '一' in s: quarter = "Q1"
        elif '二' in s: quarter = "Q2"
        elif '三' in s: quarter = "Q3"
        elif '四' in s: quarter = "Q4"
        else:
            rem = s.replace(year, '')
            if '1' in rem: quarter = "Q1"
            elif '2' in rem: quarter = "Q2"
            elif '3' in rem: quarter = "Q3"
            elif '4' in rem: quarter = "Q4"
        return f"{year}{quarter}"

    @staticmethod
    def get_default_config():
        return pd.DataFrame([
            {"所属表": "结果表3", "序号": 1, "目标字段": "序号", "源表": "🔒 系统生成", "匹配字段": "-", "逻辑说明": "自增序列"},
            {"所属表": "结果表3", "序号": 2, "目标字段": "人员", "源表": "Source A", "匹配字段": "人员", "逻辑说明": "主键 (Join Key)"},
            {"所属表": "结果表3", "序号": 3, "目标字段": "所属项目", "源表": "Source A", "匹配字段": "项目", "逻辑说明": "维度 (First)"},
            {"所属表": "结果表3", "序号": 4, "目标字段": "人事范围", "源表": "Source A", "匹配字段": "人事范围", "逻辑说明": "维度 (->销售公司)"},
            {"所属表": "结果表3", "序号": 5, "目标字段": "SPM", "源表": "Source A", "匹配字段": "SPM", "逻辑说明": "主键 (Join Key)"},
            {"所属表": "结果表3", "序号": 6, "目标字段": "合同主体", "源表": "Source A", "匹配字段": "合同主体", "逻辑说明": "维度 (->采购公司)"},
            {"所属表": "结果表3", "序号": 7, "目标字段": "销售人员", "源表": "Source A", "匹配字段": "销售", "逻辑说明": "维度 (First)"},
            {"所属表": "结果表3", "序号": 8, "目标字段": "销售部门", "源表": "Source A", "匹配字段": "销售部门", "逻辑说明": "维度 (->采购部门)"},
            {"所属表": "结果表3", "序号": 9, "目标字段": "差旅补助", "源表": "Source B", "匹配字段": "金额", "逻辑说明": "筛选：类型='差旅补助' | 清洗：x10000转元"},
            {"所属表": "结果表3", "序号": 10, "目标字段": "差旅费控平台", "源表": "Source B", "匹配字段": "金额", "逻辑说明": "筛选：类型!='差旅补助' | 清洗：x10000转元"},
            {"所属表": "结果表3", "序号": 11, "目标字段": "耗时（小时）", "源表": "Source A", "匹配字段": "交付工时", "逻辑说明": "SUM聚合 (清洗去逗号)"},
            {"所属表": "结果表3", "序号": 12, "目标字段": "支持时间（人天）", "源表": "🔒 公式计算", "匹配字段": "-", "逻辑说明": "耗时 / 8"},
            {"所属表": "结果表3", "序号": 13, "目标字段": "人力费用", "源表": "🔒 公式计算", "匹配字段": "-", "逻辑说明": "人天 * 单价 (保留2位小数)"},
            {"所属表": "结果表3", "序号": 14, "目标字段": "结算费用合计", "源表": "🔒 公式计算", "匹配字段": "-", "逻辑说明": "人力 + 差旅 + 费控 (保留2位小数)"},
            {"所属表": "结果表3", "序号": 15, "目标字段": "[配置] B表关联人", "源表": "Source B", "匹配字段": "出差人", "逻辑说明": "辅助：用于匹配A表人员 (自动去'_云计算'后缀)"},
            {"所属表": "结果表3", "序号": 16, "目标字段": "[配置] B表关联SPM", "源表": "Source B", "匹配字段": "SPM", "逻辑说明": "辅助：用于匹配A表SPM"},
            {"所属表": "结果表3", "序号": 17, "目标字段": "[配置] B表类型列", "源表": "Source B", "匹配字段": "产品类型", "逻辑说明": "辅助：用于区分补助/费控"},
            {"所属表": "结果表2", "序号": 1, "目标字段": "序号", "源表": "🔒 系统生成", "匹配字段": "-", "逻辑说明": "自增序列"},
            {"所属表": "结果表2", "序号": 2, "目标字段": "销售公司", "源表": "🔒 结果表3", "匹配字段": "人事范围", "逻辑说明": "维度分组"},
            {"所属表": "结果表2", "序号": 3, "目标字段": "采购公司", "源表": "🔒 结果表3", "匹配字段": "合同主体", "逻辑说明": "维度分组"},
            {"所属表": "结果表2", "序号": 4, "目标字段": "采购部门", "源表": "🔒 结果表3", "匹配字段": "销售部门", "逻辑说明": "维度分组"},
            {"所属表": "结果表2", "序号": 5, "目标字段": "金额（含税）", "源表": "🔒 结果表3", "匹配字段": "结算费用合计", "逻辑说明": "SUM聚合 (保留2位小数)"},
            {"所属表": "结果表2", "序号": 6, "目标字段": "工作量（人天）", "源表": "🔒 结果表3", "匹配字段": "支持时间（人天）", "逻辑说明": "SUM聚合"},
            {"所属表": "结果表1", "序号": 1, "目标字段": "序号", "源表": "🔒 系统生成", "匹配字段": "-", "逻辑说明": "自增序列"},
            {"所属表": "结果表1", "序号": 2, "目标字段": "人员", "源表": "🔒 结果表3", "匹配字段": "人员", "逻辑说明": "维度分组"},
            {"所属表": "结果表1", "序号": 3, "目标字段": "项目工时", "源表": "🔒 结果表3", "匹配字段": "耗时（小时）", "逻辑说明": "SUM聚合"},
        ])

    @staticmethod
    def get_col(config_df, target):
        row = config_df[config_df['目标字段'] == target]
        if row.empty: return None
        return str(row.iloc[0]['匹配字段']).strip()

    @staticmethod
    def clean_num(df, col):
        if col not in df.columns: return 0
        return pd.to_numeric(df[col].astype(str).str.replace(',', ''), errors='coerce').fillna(0)

    @staticmethod
    def smart_slot_check(df, slot_type):
        cols = "".join(list(df.columns))
        is_ok = True; msg = ""
        if slot_type == 'A':
            if not any(k in cols for k in ['工', '时']) and any(k in cols for k in ['金', '额', '税']):
                is_ok = False; msg = "⚠️ 警告：您似乎在 [Source A 工时] 槽位上传了 [费用表]？"
        elif slot_type == 'B':
            if not any(k in cols for k in ['金', '额', '费']) and any(k in cols for k in ['工', '时']):
                is_ok = False; msg = "⚠️ 警告：您似乎在 [Source B 差旅] 槽位上传了 [工时表]？"
        return is_ok, msg

    @staticmethod
    def validate(df_a, df_b, config_df, min_hours):
        errors = []
        c = lambda t: DataEngine.get_col(config_df, t)
        col_a_user = c('人员'); col_a_spm = c('SPM'); col_a_hrs = c('耗时（小时）')
        col_b_user = c('[配置] B表关联人'); col_b_spm = c('[配置] B表关联SPM'); col_b_amt = c('差旅补助')
        
        def check(df, col, src, tag):
            if col and col not in df.columns:
                errors.append({'类型':'逻辑错误', '来源':src, '_sys_id':'-', '行号':'-', '信息':f'缺列: {col} (用途:{tag})'})
                return False
            return True

        valid_a = True
        if df_a is not None: valid_a = check(df_a, col_a_user, 'Source A', '人员') and check(df_a, col_a_spm, 'Source A', 'SPM') and check(df_a, col_a_hrs, 'Source A', '工时')
        valid_b = True
        if df_b is not None: valid_b = check(df_b, col_b_user, 'Source B', '出差人') and check(df_b, col_b_spm, 'Source B', 'SPM') and check(df_b, col_b_amt, 'Source B', '金额')

        if not (valid_a and valid_b): return errors, df_a, df_b

        if df_a is not None and df_b is not None:
            if col_b_user in df_b.columns and col_a_user in df_a.columns:
                users_in_a = set(df_a[col_a_user].astype(str).str.strip())
                b_clean_series = df_b[col_b_user].astype(str).str.replace('_云计算', '', regex=False).str.strip()
                users_in_b = set(b_clean_series)
                ghost_users = users_in_b - users_in_a
                if ghost_users:
                    for u in ghost_users:
                        example_rows = df_b[b_clean_series == u]
                        for idx, row in example_rows.iterrows():
                            sys_id = row.get('_sys_id', idx+1)
                            errors.append({'类型': '业务规则校验', '来源': 'Source B', '_sys_id': sys_id, '行号': sys_id, '信息': f'异常差旅：人员【{u}】产生差旅费用，但在 Source A 中无对应交付记录'})

        if df_a is not None:
            df_a_clean = df_a.copy()
            df_a_clean[col_a_hrs] = DataEngine.clean_num(df_a_clean, col_a_hrs)
            for i, r in df_a_clean[df_a_clean[col_a_hrs] < 0].iterrows():
                errors.append({'类型':'数据错误', '来源':'Source A', '_sys_id':r.get('_sys_id','-'), '行号':r.get('_sys_id','-'), '信息':'工时为负'})
            if col_a_user and col_a_hrs and min_hours > 0:
                user_sums = df_a_clean.groupby(col_a_user)[col_a_hrs].sum()
                invalid_users = user_sums[user_sums < min_hours]
                for user, total_hrs in invalid_users.items():
                    sample_rows = df_a_clean[df_a_clean[col_a_user] == user]
                    if not sample_rows.empty:
                        r = sample_rows.iloc[0]
                        errors.append({'类型': '业务规则校验', '来源': 'Source A', '_sys_id': r.get('_sys_id', '-'), '行号': r.get('_sys_id', '-'), '信息': f'人员【{user}】总工时({total_hrs}h) 低于阈值({min_hours}h)'})
        return errors, df_a, df_b

    @staticmethod
    def calculate(df_a, df_b, config_df, price_per_day, subsidy_tag):
        if df_a is None or df_b is None: return None
        c = lambda t: DataEngine.get_col(config_df, t)
        col_a_user = c('人员'); col_a_spm = c('SPM'); col_a_hrs = c('耗时（小时）')
        dims_a = {'project': c('所属项目'), 'range': c('人事范围'), 'contract': c('合同主体'), 'sales': c('销售人员'), 'dept': c('销售部门')}
        col_b_user = c('[配置] B表关联人'); col_b_spm = c('[配置] B表关联SPM'); col_b_amt = c('差旅补助'); col_b_type = c('[配置] B表类型列')

        df_a[col_a_hrs] = DataEngine.clean_num(df_a, col_a_hrs)
        df_b[col_b_amt] = DataEngine.clean_num(df_b, col_b_amt)
        df_b[col_b_amt] = df_b[col_b_amt].round(2)
        if col_b_user and col_b_user in df_b.columns: df_b[col_b_user] = df_b[col_b_user].astype(str).str.replace('_云计算', '', regex=False).str.strip()

        agg_rules = {col_a_hrs: 'sum'}
        for _, col in dims_a.items(): 
            if col: agg_rules[col] = 'first'
        df_a_gp = df_a.groupby([col_a_user, col_a_spm], as_index=False).agg(agg_rules)

        is_sub = df_b[col_b_type].astype(str).str.contains(subsidy_tag, na=False)
        grp_b = [col_b_user, col_b_spm]
        df_sub = df_b[is_sub].groupby(grp_b)[col_b_amt].sum().reset_index(name='差旅补助')
        df_fee = df_b[~is_sub].groupby(grp_b)[col_b_amt].sum().reset_index(name='差旅费控平台')

        for d in [df_a_gp, df_sub, df_fee]:
            k = col_a_spm if col_a_spm in d.columns else col_b_spm
            d[k] = d[k].astype(str).str.strip()

        res = pd.merge(df_a_gp, df_sub, left_on=[col_a_user, col_a_spm], right_on=[col_b_user, col_b_spm], how='left')
        res = pd.merge(res, df_fee, left_on=[col_a_user, col_a_spm], right_on=[col_b_user, col_b_spm], how='left')
        res = res.fillna(0)

        res['支持时间（人天）'] = res[col_a_hrs] / 8
        res['人力费用'] = (res['支持时间（人天）'] * price_per_day).round(2)
        res['差旅补助'] = res['差旅补助'].round(2)
        res['差旅费控平台'] = res['差旅费控平台'].round(2)
        res['结算费用合计'] = (res['人力费用'] + res['差旅补助'] + res['差旅费控平台']).round(2)

        rename_map = {col_a_user: '人员', dims_a['project']: '所属项目', dims_a['range']: '人事范围', col_a_spm: 'SPM', dims_a['contract']: '合同主体', dims_a['sales']: '销售人员', dims_a['dept']: '销售部门', col_a_hrs: '耗时（小时）'}
        t3 = res.rename(columns=rename_map)
        final_cols = ['序号','人员','所属项目','人事范围','SPM','合同主体','销售人员','销售部门','差旅补助','差旅费控平台','耗时（小时）','支持时间（人天）','人力费用','结算费用合计']
        t3.insert(0, '序号', range(1, len(t3)+1))
        t3 = t3[[c for c in final_cols if c in t3.columns]]

        t2_cols = ['人事范围', '合同主体', '销售部门']
        if all(c in t3.columns for c in t2_cols):
            t2 = t3.groupby(t2_cols).agg({'结算费用合计':'sum', '支持时间（人天）':'sum'}).reset_index()
            t2['结算费用合计'] = t2['结算费用合计'].round(2)
            t2.columns = ['销售公司', '采购公司', '采购部门', '金额（含税，单位：元）', '工作量（人天）']
            t2.insert(0, '序号', range(1, len(t2)+1))
        else: t2 = pd.DataFrame()

        t1 = t3.groupby('人员')['耗时（小时）'].sum().reset_index()
        t1.rename(columns={'耗时（小时）':'项目工时'}, inplace=True)
        t1['人员类型'] = t1['人员'].apply(lambda x: '实施交付部' if str(x).strip() == '黄毅兵' else '实施交付部云交付小组')
        t1['备注'] = ''
        t1.insert(0, '序号', range(1, len(t1)+1))
        t1 = t1[['序号', '人员', '人员类型', '项目工时', '备注']]
        return {'t1': t1, 't2': t2, 't3': t3}

    @staticmethod
    def verify_balance(df_a, df_b, results_dict, config_df):
        if results_dict is None: return True, ""
        messages = []
        is_balanced = True
        c = lambda t: DataEngine.get_col(config_df, t)
        col_a_hrs = c('耗时（小时）'); col_b_amt = c('差旅补助')
        df_t1 = results_dict['t1']; df_t2 = results_dict['t2']; df_t3 = results_dict['t3']
        
        clean_a_hrs = DataEngine.clean_num(df_a, col_a_hrs).sum()
        res_hrs = df_t3['耗时（小时）'].sum()
        if abs(clean_a_hrs - res_hrs) > 0.1:
            is_balanced = False; messages.append(f"❌ [输入输出] 工时丢失：源表({clean_a_hrs:,.1f}) != 明细表({res_hrs:,.1f})")
            
        clean_b_amt = DataEngine.clean_num(df_b, col_b_amt).sum()
        res_amt = df_t3['差旅补助'].sum() + df_t3['差旅费控平台'].sum()
        if abs(clean_b_amt - res_amt) > 0.1:
            is_balanced = False; messages.append(f"❌ [输入输出] 金额丢失：源表({clean_b_amt:,.2f}) != 明细表({res_amt:,.2f})")

        req_cols = ['合同主体', '人事范围', '销售部门']
        if all(c in df_t3.columns for c in req_cols):
            df_t4 = df_t3.groupby(req_cols)[['结算费用合计', '支持时间（人天）']].sum().reset_index()
            t2_sum_amt = df_t2['金额（含税，单位：元）'].sum(); t4_sum_amt = df_t4['结算费用合计'].sum()
            if abs(t2_sum_amt - t4_sum_amt) > 0.05:
                is_balanced = False; messages.append(f"❌ [内部勾稽] 结算汇总表(T2)与分单合集(T4)金额不平")
            t2_sum_days = df_t2['工作量（人天）'].sum(); t4_sum_days = df_t4['支持时间（人天）'].sum()
            if abs(t2_sum_days - t4_sum_days) > 0.05:
                is_balanced = False; messages.append(f"❌ [内部勾稽] 结算汇总表(T2)与分单合集(T4)人天不平")
            t1_sum_hrs = df_t1['项目工时'].sum(); t4_calc_hrs = df_t4['支持时间（人天）'].sum() * 8
            if abs(t1_sum_hrs - t4_calc_hrs) > 0.1:
                is_balanced = False; messages.append(f"❌ [内部勾稽] 工时统计表(T1)与分单合集(T4)工时转换不平")
            t3_sum_amt = df_t3['结算费用合计'].sum()
            if abs(t3_sum_amt - t4_sum_amt) > 0.05:
                 is_balanced = False; messages.append(f"❌ [内部勾稽] 明细底表(T3)与分单合集(T4)金额聚合不平")

        if is_balanced: return True, "✅ 全链路校验通过：输入输出平衡，且 Result 1/2/3/4 内部勾稽完全一致。"
        else: return False, " | ".join(messages)

    @staticmethod
    def to_bytes(df, title=None):
        b = io.BytesIO()
        out = df.drop(columns=['_sys_id'], errors='ignore')
        start_row = 1 if title else 0
        with pd.ExcelWriter(b, engine='openpyxl') as writer:
            out.to_excel(writer, index=False, sheet_name='Sheet1', startrow=start_row)
            worksheet = writer.sheets['Sheet1']
            thin = Side(border_style="thin", color="000000"); border = Border(top=thin, left=thin, right=thin, bottom=thin)
            align_center = Alignment(horizontal='center', vertical='center', wrap_text=False); header_font = Font(bold=True)
            max_r = len(out) + 1 + start_row; max_c = len(out.columns)
            for row in worksheet.iter_rows(min_row=1, max_row=max_r, min_col=1, max_col=max_c):
                for cell in row:
                    cell.border = border; cell.alignment = align_center
                    if cell.row == (start_row + 1): cell.font = header_font
            if title:
                worksheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max_c)
                title_cell = worksheet.cell(row=1, column=1); title_cell.value = title
                title_cell.font = Font(name='SimSun', bold=True, size=18)
                title_cell.alignment = Alignment(horizontal='center', vertical='center')
                worksheet.row_dimensions[1].height = 30
            for i, col in enumerate(out.columns):
                max_len = 0
                try: max_len = len(str(col).encode('gbk'))
                except: max_len = len(str(col))
                for val in out[col]:
                    v_len = 0
                    try: v_len = len(str(val).encode('gbk'))
                    except: v_len = len(str(val))
                    if v_len > max_len: max_len = v_len
                adjusted_width = min((max_len + 2) * 1.1, 60) 
                worksheet.column_dimensions[get_column_letter(i + 1)].width = adjusted_width
        return b.getvalue()

class WordGenerator:
    @staticmethod
    def set_cell_style(cell, text, font_size=10, bold=False, align="center"):
        cell.text = ""
        paragraph = cell.paragraphs[0]
        if align == "center": paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "left": paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif align == "right": paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run(str(text))
        run.font.bold = bold; run.font.size = Pt(font_size)
        try: run.font.name = 'SimSun'; run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        except: pass
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    @staticmethod
    def set_row_height(row, height_cm):
        tr = row._tr; trPr = tr.get_or_add_trPr(); trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(int(height_cm * 567))); trHeight.set(qn('w:hRule'), "atLeast"); trPr.append(trHeight)

    @staticmethod
    def _create_base_doc(purchase_comp, sales_comp, dept_name, period_text):
        doc = docx.Document()
        section = doc.sections[0]
        section.top_margin = Cm(2.0); section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0); section.right_margin = Cm(2.0)
        title_line_1 = f"{purchase_comp}与云软件事业部-实施交付部"
        title_line_2 = f"{period_text}项目交付与运维费用结算账单"
        p1 = doc.add_paragraph(); p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p1.add_run(title_line_1); run1.font.bold = True; run1.font.size = Pt(14)
        try: run1.font.name = 'SimSun'; run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        except: pass
        p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(title_line_2); run2.font.size = Pt(14)
        try: run2.font.name = 'SimSun'; run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        except: pass
        doc.add_paragraph() 
        table0 = doc.add_table(rows=6, cols=10); table0.style = 'Table Grid'
        col_widths = [1.3, 1.3, 1.3, 1.6, 1.6, 1.6, 1.6, 1.8, 1.3, 2.0]
        for row in table0.rows:
            for idx, width in enumerate(col_widths): row.cells[idx].width = Cm(width)
        WordGenerator.set_row_height(table0.rows[0], 1.63)
        c0 = table0.rows[0].cells[0].merge(table0.rows[0].cells[9])
        WordGenerator.set_cell_style(c0, title_line_1 + "\n" + title_line_2, font_size=10, bold=True)
        WordGenerator.set_row_height(table0.rows[1], 0.92)
        WordGenerator.set_cell_style(table0.rows[1].cells[0].merge(table0.rows[1].cells[2]), "工作量\n(单位: 人/天)")
        WordGenerator.set_cell_style(table0.rows[1].cells[3].merge(table0.rows[1].cells[5]), "人力费用\n(单位: 元)")
        WordGenerator.set_cell_style(table0.rows[1].cells[6].merge(table0.rows[1].cells[7]), "差旅费用\n(单位: 元)")
        WordGenerator.set_cell_style(table0.rows[1].cells[8].merge(table0.rows[1].cells[9]), "合计")
        WordGenerator.set_row_height(table0.rows[2], 0.92)
        headers = [(0,"项目标\n准交付"), (1,"项目数\n据治理"), (2,"项目运\n维服务"), (3,"项目标\n准交付"), (4,"项目数\n据治理"), (5,"项目运\n维服务"), (6,"差旅\n补助"), (7,"商旅平\n台费用"), (8,"工作量"), (9,"合计费用\n（单位:元）")]
        for idx, txt in headers: WordGenerator.set_cell_style(table0.rows[2].cells[idx], txt)
        WordGenerator.set_row_height(table0.rows[3], 0.92)
        WordGenerator.set_row_height(table0.rows[4], 1.13)
        WordGenerator.set_cell_style(table0.rows[4].cells[0].merge(table0.rows[4].cells[2]), "项目所属区域")
        WordGenerator.set_cell_style(table0.rows[4].cells[3].merge(table0.rows[4].cells[9]), str(dept_name), align="left")
        WordGenerator.set_row_height(table0.rows[5], 4.17)
        WordGenerator.set_cell_style(table0.rows[5].cells[0].merge(table0.rows[5].cells[2]), "项目所属\n区域销售\n确认")
        c_sign = table0.rows[5].cells[3].merge(table0.rows[5].cells[9]); c_sign.text = ""
        p = c_sign.add_paragraph("确认意见：\n\n\n\n签字（签章）：\n\n\n"); p.runs[0].font.size = Pt(10)
        try: p.runs[0].font.name = 'SimSun'; p.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        except: pass
        p_date = c_sign.add_paragraph("日期：    年    月    日        "); p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        try: p_date.runs[0].font.name = 'SimSun'; p_date.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        except: pass
        doc.add_paragraph("\n")
        table1 = doc.add_table(rows=1, cols=11); table1.style = 'Table Grid'
        t1_widths = [0.9, 1.2, 1.5, 1.5, 1.0, 1.0, 1.0, 2.2, 2.3, 2.3, 2.6]
        for row in table1.rows:
            for idx, width in enumerate(t1_widths): row.cells[idx].width = Cm(width)
        headers_1 = ['人员', '人事\n范围', '项目\n名称', '合同\n名称', '销售\n人员', '销售所\n在大区', '支持\n人天', '人力\n费用', '差旅\n补助', '差旅平\n台费用', '总费用\n（元）']
        for i, text in enumerate(headers_1): WordGenerator.set_cell_style(table1.rows[0].cells[i], text)
        return doc, table0, table1

    @staticmethod
    def generate(df_t3, period_str):
        if not HAS_DOCX:
            return {}, "缺少 python-docx 库，无法生成 Word 文件"
        if df_t3 is None or df_t3.empty:
            return {}, "数据为空"
            
        req = ['合同主体', '人事范围', '销售部门']
        if not all(c in df_t3.columns for c in req):
            return {}, f"Result T3 缺少必要的列: {req}"

        files_dict = {}
        grouped = df_t3.groupby(req)

        for (purch_comp, sales_comp, dept_name), group in grouped:
            doc, table0, table1 = WordGenerator._create_base_doc(purch_comp, sales_comp, dept_name, period_str)
            
            for _, row in group.iterrows():
                cells = table1.add_row().cells
                WordGenerator.set_cell_style(cells[0], row['人员'])
                WordGenerator.set_cell_style(cells[1], row['人事范围'])
                WordGenerator.set_cell_style(cells[2], row['所属项目'])
                WordGenerator.set_cell_style(cells[3], row['合同主体'])
                WordGenerator.set_cell_style(cells[4], row['销售人员'])
                WordGenerator.set_cell_style(cells[5], row['销售部门']) 
                WordGenerator.set_cell_style(cells[6], f"{row['支持时间（人天）']:.2f}")
                WordGenerator.set_cell_style(cells[7], f"{row['人力费用']:.2f}")
                WordGenerator.set_cell_style(cells[8], f"{row['差旅补助']:.2f}")
                WordGenerator.set_cell_style(cells[9], f"{row['差旅费控平台']:.2f}")
                WordGenerator.set_cell_style(cells[10], f"{row['结算费用合计']:.2f}")

            sum_days = group['支持时间（人天）'].sum()
            sum_man_cost = group['人力费用'].sum()
            total_cost = group['结算费用合计'].sum()
            
            WordGenerator.set_cell_style(table0.rows[3].cells[0], f"{sum_days:.2f}") 
            WordGenerator.set_cell_style(table0.rows[3].cells[1], "-")
            WordGenerator.set_cell_style(table0.rows[3].cells[2], "-")
            WordGenerator.set_cell_style(table0.rows[3].cells[3], f"{sum_man_cost:.2f}") 
            WordGenerator.set_cell_style(table0.rows[3].cells[4], "-")
            WordGenerator.set_cell_style(table0.rows[3].cells[5], "-")
            WordGenerator.set_cell_style(table0.rows[3].cells[6], f"{group['差旅补助'].sum():.2f}")
            WordGenerator.set_cell_style(table0.rows[3].cells[7], f"{group['差旅费控平台'].sum():.2f}")
            WordGenerator.set_cell_style(table0.rows[3].cells[8], f"{sum_days:.2f}") 
            WordGenerator.set_cell_style(table0.rows[3].cells[9], f"{total_cost:.2f}") 

            out = io.BytesIO()
            doc.save(out)
            safe_dept = str(dept_name).replace('/', '_').replace('\\', '_')
            fname = f"结算单_{purch_comp}_{sales_comp}_{safe_dept}.docx"
            files_dict[fname] = out.getvalue()
            
        return files_dict, None

# ==============================================================================
# Zone B: UI 组件层
# ==============================================================================
class UIComponents:
    @staticmethod
    def render_sidebar(threshold_error_flag):
        with st.sidebar:
            st.markdown("<div class='nav-header'>⚙️ 参数配置</div>", unsafe_allow_html=True)
            st.divider()
            
            st.session_state.params['price'] = st.number_input("人力单价 (元/天)", value=st.session_state.params['price'], step=100)
            if threshold_error_flag: st.error("🚨 请调整工时", icon=None)
            st.session_state.params['hours_limit'] = st.number_input("工时阈值 (小时)", value=st.session_state.params['hours_limit'])
            # st.session_state.params['sub_tag'] 已在 init_defaults 中硬编码，无需显示输入框
            st.session_state.params['period'] = st.text_input("结算周期文案", value=st.session_state.params['period'])
            
            current_params = st.session_state.params.copy()
            last_run_params = st.session_state.get('last_run_params', None)
            has_files = st.session_state.data_store['A']['df'] is not None and st.session_state.data_store['B']['df'] is not None
            param_changed = last_run_params is not None and current_params != last_run_params
            
            trigger_recalc = False
            if has_files and param_changed:
                st.write("") 
                if st.button("重新运算", type="primary", use_container_width=True): trigger_recalc = True
            
            # ====== 核心修复逻辑：Flex 占位符 ======
            # 利用父容器的 Flex 属性，这个空的 div 会自动撑开所有剩余空间
            st.markdown('<div style="flex: 1;"></div>', unsafe_allow_html=True)
            
            # 猫咪按钮
            if st.button("🐱", key="btn_cat_config", type="tertiary", help="进入规则配置中心"):
                st.session_state.page = 'mapping'; st.session_state.prank_solved = False; st.rerun()
            
            # 底部留白
            st.markdown('<div style="height: 15px;"></div>', unsafe_allow_html=True)

            return current_params, trigger_recalc

    @staticmethod
    @st.cache_data(ttl=3600)
    def load_data_cached(file_content, file_name):
        try:
            if file_name.endswith('.csv'): df = pd.read_csv(file_content)
            else: df = pd.read_excel(file_content, engine='openpyxl')
            df.columns = [str(c).strip() for c in df.columns]; df['_sys_id'] = range(1, len(df)+1)
            return df
        except Exception: return None

    @staticmethod
    def render_file_slot(key, title, data_store):
        data = data_store[key]
        has_file = data['df'] is not None
        with st.container(border=True):
            if not has_file:
                st.markdown(f"**{title}**")
                file = st.file_uploader(title, type=['xlsx', 'csv'], key=f"uploader_{key}", label_visibility="collapsed")
                if file:
                    df = UIComponents.load_data_cached(file, file.name)
                    if df is not None:
                        is_ok, msg = DataEngine.smart_slot_check(df, key)
                        if not is_ok: st.toast(msg, icon="🚨"); time.sleep(2)
                        st.session_state.data_store[key] = {'df': df, 'name': file.name}
                        st.session_state.is_calculated = False; st.session_state.error_report = None; st.session_state.balance_check = (True, ""); st.session_state.last_run_params = None; st.rerun()
            else:
                c_icon, c_info, c_close = st.columns([0.15, 0.70, 0.15], vertical_alignment="center")
                with c_icon: st.markdown('<div class="file-icon">📄</div>', unsafe_allow_html=True)
                with c_info:
                    row_str = "{:,}".format(len(data['df'])); st.markdown(f"""<div style="display: flex; flex-direction: column; justify-content: center; height: 100%;"><span class="file-name">{data['name']}</span><span class="file-stats">📊 已加载 {row_str} 条数据</span></div>""", unsafe_allow_html=True)
                with c_close:
                    if st.button("✕", key=f"del_{key}", help="移除此文件", type="secondary"): 
                        st.session_state.data_store[key] = {'df': None, 'name': None}
                        st.session_state.is_calculated = False; st.session_state.error_report = None; st.session_state.balance_check = (True, ""); st.session_state.last_run_params = None; st.rerun()

    @staticmethod
    def render_error_report(err_df):
        fixable = err_df[err_df['类型']=='数据错误']; rule = err_df[err_df['类型']=='业务规则校验']
        st.markdown(f"<div class='error-box'><h3 style='margin:0'>🚨 校验未通过</h3><p>发现 <b>{len(fixable)}</b> 个数据错误，<b>{len(rule)}</b> 个业务规则警告。</p></div>", unsafe_allow_html=True)
        st.dataframe(err_df[['类型','来源','行号','信息']], use_container_width=True, hide_index=True)
        st.download_button("📥 下载错误清单", err_df.to_csv(index=False).encode('utf-8-sig'), "err.csv", "text/csv", use_container_width=True)
        return any("阈值" in str(x) for x in err_df['信息'])

    @staticmethod
    def render_download_zone(result_files, all_in_one_zip, word_files_dict, period_str, balance_check):
        is_bal, bal_msg = balance_check; css_class = "balance-box-ok" if is_bal else "balance-box-err"
        st.markdown(f"<div class='{css_class}'>{bal_msg}</div>", unsafe_allow_html=True)
        if not is_bal: st.warning("⚠️ 严重警告：总额或内部勾稽不平，请务必检查上方错误信息！")
        with st.container(border=True):
            st.success("✅ 计算完成 | 报表已生成"); st.subheader("📦 批量下载")
            st.download_button("🚀 一键下载 (Excel + Word)", all_in_one_zip, "项目结算资料全集.zip", "application/zip", type="primary", use_container_width=True)
            st.divider(); q_str = DataEngine.get_quarter_str(period_str)
            name_t1 = f"实施交付部项目情况汇总_部门工时统计-{q_str}.xlsx"; name_t3 = f"实施交付部项目情况汇总_结算工时总表-{q_str}.xlsx"; name_t2 = f"{q_str}实施交付部项目投入考核调整总表.xlsx"
            c1, c2, c3 = st.columns(3)
            if result_files and 't1' in result_files: c1.download_button(f"📥 {name_t1}", result_files['t1'], name_t1, use_container_width=True)
            if result_files and 't2' in result_files: c2.download_button(f"📥 {name_t2}", result_files['t2'], name_t2, use_container_width=True)
            if result_files and 't3' in result_files: c3.download_button(f"📥 {name_t3}", result_files['t3'], name_t3, use_container_width=True)
            if word_files_dict:
                with st.expander(f"📄 结算单 ({len(word_files_dict)})", expanded=False):
                    for fname, fbytes in word_files_dict.items():
                        c_t, c_b = st.columns([4, 1]); c_t.text(f"📄 {fname}")
                        c_b.download_button("下载", fbytes, fname, key=f"btn_{fname}")

    @staticmethod
    def render_native_editor(desc, subset, is_readonly, all_options):
        if subset.empty: return None
        
        column_config = {
            "序号": st.column_config.TextColumn("序号", width="small", disabled=True),
            "目标字段": st.column_config.TextColumn("目标字段", disabled=True, width="medium"),
            "源表": st.column_config.TextColumn("源表", disabled=True),
            "逻辑说明": st.column_config.TextColumn("逻辑说明", disabled=True, width="large"),
        }
        
        if is_readonly:
            column_config["匹配字段"] = st.column_config.TextColumn("匹配字段", disabled=True)
        else:
            column_config["匹配字段"] = st.column_config.SelectboxColumn(
                "匹配字段", options=all_options, width="medium", required=True
            )
        
        calc_height = (len(subset) + 1) * 35 + 10
        editor_key = f"editor_{desc}_{subset.iloc[0]['所属表']}_{st.session_state.editing_template_name}"
        st.markdown(f"**{desc}**")
        return st.data_editor(subset, column_config=column_config, use_container_width=True, hide_index=True, height=max(150, min(1000, calc_height)), key=editor_key, disabled=is_readonly)

# ==============================================================================
# Zone C: 控制层
# ==============================================================================
if 'page' not in st.session_state: st.session_state.page = 'main'
if 'data_store' not in st.session_state: st.session_state.data_store = {'A': {'df': None, 'name': None}, 'B': {'df': None, 'name': None}}
if 'is_calculated' not in st.session_state: st.session_state.is_calculated = False
if 'error_report' not in st.session_state: st.session_state.error_report = None
if 'all_zip' not in st.session_state: st.session_state.all_zip = None
if 'word_files' not in st.session_state: st.session_state.word_files = {}
if 'result_files' not in st.session_state: st.session_state.result_files = {}
if 'last_run_params' not in st.session_state: st.session_state.last_run_params = None
if 'threshold_error_flag' not in st.session_state: st.session_state.threshold_error_flag = False
if 'balance_check' not in st.session_state: st.session_state.balance_check = (True, "")
if 'sample_store' not in st.session_state: st.session_state.sample_store = {'A': None, 'B': None}

TemplateManager.init_defaults()
inject_css()

if st.session_state.page == 'main':
    current_params, manual_recalc = UIComponents.render_sidebar(st.session_state.threshold_error_flag)
    st.title("😈 淡藤财务报表 Pro")
    with st.container(border=True):
        st.markdown("### 📂 数据源控制台")
        st.divider()
        
        # --- Template Selector ---
        all_templates = TemplateManager.get_all_names()
        if st.session_state.active_template_name not in all_templates: 
            st.session_state.active_template_name = TemplateManager.DEFAULT_NAME
        
        c_sel, c_empty = st.columns([1.5, 8.5])
        with c_sel:
            selected_tpl = st.selectbox(
                "计算规则模板", 
                options=all_templates, 
                index=all_templates.index(st.session_state.active_template_name), 
                key="main_template_selector", 
                label_visibility="collapsed"
            )
            
        if selected_tpl != st.session_state.active_template_name:
            st.session_state.active_template_name = selected_tpl; st.session_state.is_calculated = False; st.rerun()
        # -------------------------

        c1, c2 = st.columns(2)
        with c1: UIComponents.render_file_slot('A', "Source A: 投入明细 (工时)", st.session_state.data_store)
        with c2: UIComponents.render_file_slot('B', "Source B: 差旅明细 (费用)", st.session_state.data_store)
        
        if st.button("🗑️ 清空所有文件", type="secondary", use_container_width=True): 
            st.session_state.data_store = {'A': {'df': None, 'name': None}, 'B': {'df': None, 'name': None}}
            st.session_state.is_calculated = False; st.session_state.error_report = None; st.session_state.all_zip = None; st.session_state.last_run_params = None; st.rerun()
        
    st.divider()
    
    has_files = st.session_state.data_store['A']['df'] is not None and st.session_state.data_store['B']['df'] is not None
    should_calculate = False
    if has_files:
        if st.session_state.last_run_params is None: should_calculate = True
        elif manual_recalc: should_calculate = True
        
        if should_calculate:
            active_config = TemplateManager.get_template(st.session_state.active_template_name)
            with st.spinner(f"🚀 正在使用 [{st.session_state.active_template_name}] 模板计算..."):
                st.session_state.threshold_error_flag = False
                errs, df_a, df_b = DataEngine.validate(
                    st.session_state.data_store['A']['df'].copy(), st.session_state.data_store['B']['df'].copy(), active_config, current_params['hours_limit']
                )
                if errs:
                    st.session_state.error_report = pd.DataFrame(errs); st.session_state.is_calculated = False; st.session_state.last_run_params = current_params.copy(); st.rerun()
                else:
                    res = DataEngine.calculate(df_a, df_b, active_config, current_params['price'], current_params['sub_tag'])
                    if res:
                        st.session_state.balance_check = DataEngine.verify_balance(df_a, df_b, res, active_config)
                        q_str = DataEngine.get_quarter_str(current_params['period'])
                        t2_title = f"{q_str[2:]}实施交付部项目投入考核调整总表"
                        excel_files_dict = {
                            "t1": DataEngine.to_bytes(res['t1']),
                            "t2": DataEngine.to_bytes(res['t2'], title=t2_title),
                            "t3": DataEngine.to_bytes(res['t3'])
                        }
                        st.session_state.result_files = excel_files_dict
                        word_files_dict, err_msg = WordGenerator.generate(res['t3'], current_params['period'])
                        if err_msg: st.warning(f"Word生成受限: {err_msg}")
                        st.session_state.word_files = word_files_dict
                        all_files_to_zip = {}
                        all_files_to_zip[f"实施交付部项目情况汇总_部门工时统计-{q_str}.xlsx"] = excel_files_dict['t1']
                        all_files_to_zip[f"{q_str}实施交付部项目投入考核调整总表.xlsx"] = excel_files_dict['t2']
                        all_files_to_zip[f"实施交付部项目情况汇总_结算工时总表-{q_str}.xlsx"] = excel_files_dict['t3']
                        all_files_to_zip.update(word_files_dict)
                        buf_zip = io.BytesIO()
                        with zipfile.ZipFile(buf_zip, 'w', zipfile.ZIP_DEFLATED) as z:
                            for fname, fcontent in all_files_to_zip.items(): z.writestr(fname, fcontent)
                        st.session_state.all_zip = buf_zip.getvalue()
                        st.session_state.is_calculated = True; st.session_state.error_report = None; st.session_state.last_run_params = current_params.copy(); st.rerun()

    if st.session_state.error_report is not None:
        is_threshold_err = UIComponents.render_error_report(st.session_state.error_report)
        if is_threshold_err != st.session_state.threshold_error_flag: st.session_state.threshold_error_flag = is_threshold_err; st.rerun()
    elif st.session_state.is_calculated:
        UIComponents.render_download_zone(st.session_state.result_files, st.session_state.all_zip, st.session_state.word_files, current_params['period'], st.session_state.balance_check)

elif st.session_state.page == 'mapping':
    if 'prank_solved' not in st.session_state: st.session_state.prank_solved = False
    if not st.session_state.prank_solved:
        c1, c2 = st.columns([9, 1]); c1.write("")
        if c2.button("⬅️", key="back_from_prank", type="tertiary", use_container_width=True): st.session_state.page = 'main'; st.rerun()
        st.markdown("""<style>.prank-container { display: flex; justify-content: center; margin-top: 150px; } .prank-text { font-size: 2.5rem; color: #30363d; font-family: 'Courier New', monospace; cursor: default; } a.prank-link { text-decoration: none; color: inherit; cursor: text; } a.prank-link:hover { color: inherit; text-decoration: none; }</style><div class="prank-container"><span class="prank-text">你以为有什么<a href="?prank=1" target="_self" class="prank-link">？</a></span></div>""", unsafe_allow_html=True)
    else:
        all_templates = TemplateManager.get_all_names() 
        with st.sidebar:
            st.header("📏 规则模板管理")
            st.divider()
            st.markdown("<div class='sidebar-label'>📝 模板列表 (点击编辑)</div>", unsafe_allow_html=True)
            for tpl_name in all_templates:
                btn_type = "primary" if tpl_name == st.session_state.editing_template_name else "secondary"
                if st.button(tpl_name, key=f"btn_edit_{tpl_name}", type=btn_type, use_container_width=True):
                    st.session_state.editing_template_name = tpl_name; st.session_state.sample_store = {'A': None, 'B': None}; st.rerun()
            st.divider()
            
            with st.expander("➕ 新建模板", expanded=False):
                new_tpl_name = st.text_input("模板名称", placeholder="1-8个字符", max_chars=8)
                if st.button("创建", key="create_new_tpl", use_container_width=True):
                    if new_tpl_name and new_tpl_name not in st.session_state.templates:
                        if 1 <= len(new_tpl_name) <= 8:
                            TemplateManager.save_template(new_tpl_name, DataEngine.get_default_config().copy())
                            st.session_state.editing_template_name = new_tpl_name; st.session_state.sample_store = {'A': None, 'B': None}
                            st.success(f"模板 {new_tpl_name} 已创建"); time.sleep(0.5); st.rerun()
                        else: st.error("长度需在1-8字符之间")
                    elif new_tpl_name: st.error("名称已存在")

        c1, c2 = st.columns([8, 2], vertical_alignment="center")
        
        is_default = (st.session_state.editing_template_name == TemplateManager.DEFAULT_NAME)
        if not is_default:
            c1.markdown(f"<div class='nav-header'>📏 正在编辑: {st.session_state.editing_template_name}</div>", unsafe_allow_html=True)
        
        if c2.button("⬅️ 返回主页", type="tertiary", use_container_width=True): st.session_state.page = 'main'; st.rerun()
        st.markdown("<hr style='margin-top:0; border-color:#30363d;'>", unsafe_allow_html=True)
        
        cols_a = []; cols_b = []
        
        if not is_default:
            with st.expander("📂 上传样例数据 (用于提取列名，不参与计算)", expanded=True):
                st.caption("提示：此处上传的文件仅用于获取表头。")
                sc1, sc2 = st.columns(2)
                sample_a = sc1.file_uploader("Source A 工时统计", type=['xlsx', 'csv'], key="sample_a")
                if sample_a: 
                    df = UIComponents.load_data_cached(sample_a, sample_a.name)
                    if df is not None: 
                        is_ok, msg = DataEngine.smart_slot_check(df, 'A')
                        if not is_ok: st.toast(msg, icon="⚠️")
                        st.session_state.sample_store['A'] = list(df.columns)
                sample_b = sc2.file_uploader("Source B 差旅明细", type=['xlsx', 'csv'], key="sample_b")
                if sample_b:
                    df = UIComponents.load_data_cached(sample_b, sample_b.name)
                    if df is not None: 
                        is_ok, msg = DataEngine.smart_slot_check(df, 'B')
                        if not is_ok: st.toast(msg, icon="⚠️")
                        st.session_state.sample_store['B'] = list(df.columns)
            cols_a = st.session_state.sample_store['A'] or []
            cols_b = st.session_state.sample_store['B'] or []
        else:
            st.info("🔒 系统默认模板为只读模式，无法修改配置或上传样例。如需修改，请先创建新模板。")

        df_c = st.session_state.templates[st.session_state.editing_template_name]
        
        def save_and_validate(edited_df):
            if is_default: return
            current_config = st.session_state.templates[st.session_state.editing_template_name]
            
            for idx, row in edited_df.iterrows():
                target = row['目标字段']
                new_match = str(row['匹配字段']).strip()
                
                mask = (current_config['所属表'] == row['所属表']) & (current_config['目标字段'] == target)
                if not mask.any(): continue
                
                orig_idx = current_config[mask].index[0]
                if new_match != str(current_config.at[orig_idx, '匹配字段']).strip():
                    st.session_state.templates[st.session_state.editing_template_name].at[orig_idx, '匹配字段'] = new_match

        t1, t2, t3 = st.tabs(["结果表3 (底表)", "结果表2 (结算)", "结果表1 (工时)"])
        
        current_options = []
        if cols_a: current_options.extend(cols_a)
        if cols_b: current_options.extend(cols_b)
        current_options = list(dict.fromkeys(current_options))
        
        with t1: 
            df_t3 = df_c[df_c["所属表"]=="结果表3"]
            
            st.markdown("#### 📂 Source A 字段映射 (工时统计)")
            df_a_subset = df_t3[df_t3["源表"] == "Source A"]
            if not cols_a and not is_default:
                st.warning("⚠️ 请先在上方上传 'Source A 工时统计' 样例数据以解锁此区域。")
            edited_a = UIComponents.render_native_editor(
                "Source A 配置", df_a_subset, is_default or not cols_a, cols_a
            )
            if not is_default and edited_a is not None: save_and_validate(edited_a)
            
            st.divider()
            
            st.markdown("#### 📂 Source B 字段映射 (差旅明细)")
            df_b_subset = df_t3[df_t3["源表"] == "Source B"]
            if not cols_b and not is_default:
                st.warning("⚠️ 请先在上方上传 'Source B 差旅明细' 样例数据以解锁此区域。")
            edited_b = UIComponents.render_native_editor(
                "Source B 配置", df_b_subset, is_default or not cols_b, cols_b
            )
            if not is_default and edited_b is not None: save_and_validate(edited_b)
            
            st.divider()
            
            st.markdown("#### 🔒 系统锁定/公式计算字段")
            df_lock_subset = df_t3[~df_t3["源表"].isin(["Source A", "Source B"])]
            UIComponents.render_native_editor("系统配置 (只读)", df_lock_subset, True, [])
        with t2: 
            st.info("ℹ️ 结果表 2 为衍生汇总表，规则由系统锁定。")
            UIComponents.render_native_editor("结算汇总表", df_c[df_c["所属表"]=="结果表2"], True, [])
        with t3: 
            st.info("ℹ️ 结果表 1 为衍生工时表，规则由系统锁定。")
            UIComponents.render_native_editor("工时统计表", df_c[df_c["所属表"]=="结果表1"], True, [])

        if not is_default:
            st.markdown("<div class='action-btn-zone'></div>", unsafe_allow_html=True)
            bc1, bc2, bc3 = st.columns([2, 6, 2])
            with bc1:
                if st.button("🗑️ 删除模板", type="secondary", use_container_width=True):
                    if TemplateManager.delete_template(st.session_state.editing_template_name):
                        st.success("模板已删除"); time.sleep(0.5); st.rerun()
            with bc3:
                if st.button("💾 确认生效", type="primary", use_container_width=True):
                    st.toast(f"模板 [{st.session_state.editing_template_name}] 已更新并校验通过", icon="✅")
